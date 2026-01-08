import streamlit as st
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from groq import Groq
import os
from dotenv import load_dotenv
import tempfile
import shutil
import pandas as pd
from dataclasses import dataclass, field
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse, parse_qs

# Try to import Document from langchain; if unavailable, provide a lightweight fallback
try:
    from langchain.schema import Document
except Exception:
    @dataclass
    class Document:
        page_content: str
        metadata: dict = field(default_factory=dict)

# rejection phrase used by the strict assistant
REJECTION_PHRASE = "The uploaded document does not contain this information."

# Load environment variables from .env file
load_dotenv()

# Load Groq client
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# Load embeddings
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

st.set_page_config(page_title="📄 DocChatBot", layout="wide")
st.title("📄 Strict Document Chatbot")

# Initialize session state to track current document
if "current_document" not in st.session_state:
    st.session_state.current_document = None

st.markdown("""
This is a **strict, document-based AI assistant**. It answers questions ONLY using information 
from your uploaded documents. No hallucinations, no external knowledge.

⚠️ **Document-Isolated Mode**: Each new document upload creates a fresh session. 
Previous documents are discarded. Questions are answered ONLY using the current document.
""")

# Sidebar for document upload and settings
with st.sidebar:
    st.header("📤 Upload Documents")
    uploaded_file = st.file_uploader("Upload a document (PDF, DOCX, CSV, TXT)", type=["pdf", "docx", "csv", "txt"])
    
    if uploaded_file is not None:
        if st.button("Process Document", use_container_width=True):
            with st.spinner("Processing document..."):
                # Save uploaded file temporarily
                temp_dir = tempfile.mkdtemp()
                temp_path = os.path.join(temp_dir, uploaded_file.name)
                
                with open(temp_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                
                # Load and process based on file type
                ext = os.path.splitext(uploaded_file.name)[1].lower()
                docs = []

                try:
                    if ext == ".pdf":
                        loader = PyPDFLoader(temp_path)
                        docs = loader.load()

                    elif ext == ".txt":
                        with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
                            text = f.read()
                        docs = [Document(page_content=text, metadata={"source": uploaded_file.name})]

                    elif ext == ".csv":
                        try:
                            df = pd.read_csv(temp_path)
                            # convert CSV to readable markdown-style table
                            text = df.to_markdown(index=False)
                        except Exception:
                            with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
                                text = f.read()
                        docs = [Document(page_content=text, metadata={"source": uploaded_file.name})]

                    elif ext == ".docx":
                        try:
                            from docx import Document as DocxReader
                            docx = DocxReader(temp_path)
                            paragraphs = [p.text for p in docx.paragraphs if p.text]
                            text = "\n".join(paragraphs)
                            docs = [Document(page_content=text, metadata={"source": uploaded_file.name})]
                        except Exception:
                            st.error("Install 'python-docx' to support .docx files: pip install python-docx")
                            docs = []

                    else:
                        # fallback: try PDF loader, then raw text
                        try:
                            loader = PyPDFLoader(temp_path)
                            docs = loader.load()
                        except Exception:
                            with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
                                text = f.read()
                            docs = [Document(page_content=text, metadata={"source": uploaded_file.name})]

                    # Split text into chunks
                    splitter = RecursiveCharacterTextSplitter(
                        chunk_size=500,
                        chunk_overlap=50
                    )
                    chunks = splitter.split_documents(docs)

                    # DELETE old index and create fresh one (document-isolated mode)
                    import shutil as shutil_delete
                    try:
                        shutil_delete.rmtree("faiss_index")
                    except Exception:
                        pass

                    # Create new index from scratch (not appending)
                    db = FAISS.from_documents(chunks, embeddings)
                    db.save_local("faiss_index")

                    # Update session state: current document
                    st.session_state.current_document = uploaded_file.name

                    st.success(f"✅ Document '{uploaded_file.name}' loaded. Previous documents cleared. Starting fresh session.")
                finally:
                    try:
                        shutil.rmtree(temp_dir)
                    except Exception:
                        pass
    
    st.divider()
    st.header("🔗 Or Load from URL")
    url_input = st.text_input("Enter a URL (web link or YouTube video):", placeholder="https://example.com or https://youtube.com/watch?v=...")
    
    if url_input and st.button("Load from URL", use_container_width=True):
        with st.spinner("Loading content from URL..."):
            try:
                docs = []
                
                # Check if it's a YouTube URL
                if "youtube.com" in url_input or "youtu.be" in url_input:
                    try:
                        from youtube_transcript_api import YouTubeTranscriptApi
                        
                        # Extract video ID
                        if "youtu.be" in url_input:
                            video_id = url_input.split("/")[-1].split("?")[0]
                        else:
                            video_id = parse_qs(urlparse(url_input).query).get('v', [None])[0]
                        
                        if video_id:
                            # Get transcript
                            transcript = YouTubeTranscriptApi.get_transcript(video_id)
                            text = "\n".join([item['text'] for item in transcript])
                            docs = [Document(page_content=text, metadata={"source": f"YouTube: {url_input}", "type": "youtube"})]
                        else:
                            st.error("Could not extract video ID from YouTube URL")
                    except Exception as e:
                        st.error(f"Error loading YouTube transcript: {str(e)}")
                        docs = []
                else:
                    # Load web content
                    try:
                        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
                        response = requests.get(url_input, headers=headers, timeout=10)
                        response.raise_for_status()
                        
                        # Parse HTML
                        soup = BeautifulSoup(response.content, 'lxml')
                        # Remove script and style tags
                        for script in soup(["script", "style"]):
                            script.decompose()
                        
                        text = soup.get_text(separator="\n", strip=True)
                        docs = [Document(page_content=text, metadata={"source": url_input, "type": "web"})]
                    except Exception as e:
                        st.error(f"Error loading web content: {str(e)}")
                        docs = []
                
                if docs:
                    # Split text into chunks
                    splitter = RecursiveCharacterTextSplitter(
                        chunk_size=500,
                        chunk_overlap=50
                    )
                    chunks = splitter.split_documents(docs)
                    
                    # DELETE old index and create fresh one (document-isolated mode)
                    try:
                        shutil.rmtree("faiss_index")
                    except Exception:
                        pass
                    
                    # Create new index from scratch
                    db = FAISS.from_documents(chunks, embeddings)
                    db.save_local("faiss_index")
                    
                    # Update session state
                    st.session_state.current_document = url_input
                    
                    st.success(f"✅ Content loaded from URL. Starting fresh session.")
                    st.rerun()
            except Exception as e:
                st.error(f"Error processing URL: {str(e)}")
    
    st.divider()
    st.header("⚙️ Settings")
    answer_type = st.radio(
        "Answer Length Preference:",
        ["Brief Answer", "Long Answer"],
        help="Brief: 2-3 sentences | Long: Detailed explanation"
    )

# Load vector DB
try:
    db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    retriever = db.as_retriever(search_kwargs={"k": 5})
except:
    st.warning("⚠️ No documents indexed yet. Please upload a document (PDF/DOCX/CSV/TXT) in the sidebar to get started.")
    retriever = None

# Query section
if retriever is not None:
    st.divider()
    if st.session_state.current_document:
        st.header(f"❓ Ask a Question (Current: {st.session_state.current_document})")
    else:
        st.header("❓ Ask a Question")
    query = st.text_input("Enter your question about the current document:", placeholder="What information is in the document?")
    
    if query:
        with st.spinner("🔍 Searching documents and generating answer..."):
            docs = retriever.invoke(query)

            # Filter out any docs that contain the assistant's rejection phrase (avoid propagating rejections into index/context)
            filtered_docs = []
            if docs:
                for d in docs:
                    content = d.page_content if hasattr(d, "page_content") else str(d)
                    if REJECTION_PHRASE.lower() in content.lower():
                        continue
                    filtered_docs.append(d)

            if not filtered_docs:
                st.warning("❌ No relevant information found in the uploaded documents.")
            else:
                retrieved_document_chunks = "\n\n---\n\n".join([d.page_content for d in filtered_docs])

                # Strict system prompt based on user requirements
                system_prompt = f"""You are a strict, document-based AI assistant operating in DOCUMENT-ISOLATED MODE.

Your task is to answer user questions ONLY using the information present in the CURRENTLY UPLOADED document.

CURRENT DOCUMENT: {st.session_state.current_document if st.session_state.current_document else 'Unknown'}

CRITICAL RULES YOU MUST FOLLOW:

1. Use ONLY the provided document context to generate answers.
2. Do NOT use external knowledge, assumptions, or general world knowledge.
3. Do NOT reference or use information from previously uploaded documents.
4. If a question requires information from a previous document, respond with:
   "I cannot respond to the question as presented. Can you ask a different question that is semantically relevant to the provided document context?"
5. If the answer is NOT present in the CURRENT document, clearly say: "The uploaded document does not contain this information."
6. Do NOT hallucinate or guess.
7. If the user question is semantically relevant to the CURRENT document (even if exact keywords are missing), answer correctly.
8. Understand the meaning of the question, not just keywords.
9. Be factually accurate and grounded in the CURRENT document content only.

ANSWER LENGTH REQUIREMENTS:
- If user selected "Brief Answer": Provide a concise, 2–3 sentence summary.
- If user selected "Long Answer": Provide a detailed, well-explained answer using multiple sentences and structured explanation if needed.

TONE:
- Clear
- Professional
- Simple and easy to understand

Now, answer the user's question ONLY using the context from the CURRENT document provided below."""

                answer_length_instruction = (
                    "Provide a brief, 2-3 sentence answer."
                    if answer_type == "Brief Answer"
                    else "Provide a detailed, comprehensive answer with multiple sentences and structured explanation if needed."
                )

                prompt = f"""{system_prompt}

ANSWER LENGTH: {answer_length_instruction}

Context from uploaded documents:
{retrieved_document_chunks}

User Question:
{query}

Remember: Answer ONLY based on the context above. If the answer is not in the context, respond with: "The uploaded document does not contain this information.\""""

                response = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[{"role": "user", "content": prompt}]
                )
                
                answer = response.choices[0].message.content
                
                # Display answer with styling (dark background + light text for readability)
                st.divider()
                st.subheader("📝 Answer")
                st.markdown(f"""
                <div style="background-color: #0b1220; color: #e6eef8; padding: 18px; border-radius: 6px; border-left: 4px solid #1f77b4; line-height:1.6;">
                {answer}
                </div>
                """, unsafe_allow_html=True)
                
                # Show retrieved chunks (sources used for the answer)
                with st.expander("📚 Retrieved Document Chunks (sources used for answer)"):
                    st.caption("These chunks were retrieved from your uploaded documents and used as evidence to generate the answer. If the answer is not present here, the uploaded document does not contain that information.")
                    for i, doc in enumerate(filtered_docs, 1):
                        st.write(f"**Chunk {i}:**")
                        st.write(doc.page_content)
                        st.divider()

#  .\venv\Scripts\streamlit run app.py

